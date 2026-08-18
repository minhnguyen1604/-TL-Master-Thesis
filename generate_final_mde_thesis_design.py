import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_bottom_border_to_paragraph(paragraph, color_hex="888888", size="6"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_neu_mde_final_docx():
    doc = docx.Document()
    
    # 1. Page Margins (1.0 inch = 2.54 cm standard)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # 2. Running Header Configuration (Different First Page)
    section.different_first_page_header_footer = True
    
    # Running Header for pages 2 onwards (Right aligned, Italic, with GRAY bottom border)
    header = section.header
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run("Thesis Design – MDE31 – Dang Tu Linh")
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(10)
    r_head.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    r_head.italic = True
    add_bottom_border_to_paragraph(p_head, color_hex="888888", size="6")

    # Configure Styles
    styles = doc.styles

    # Normal / Body Text Style
    style_normal = styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    # Heading 1 Style
    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(14)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    style_h1.paragraph_format.space_before = Pt(14)
    style_h1.paragraph_format.space_after = Pt(6)

    # Heading 2 Style
    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(12.5)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    style_h2.paragraph_format.space_before = Pt(10)
    style_h2.paragraph_format.space_after = Pt(4)

    # ==================== COVER / TITLE PAGE ====================
    p_header = doc.add_paragraph()
    p_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_header.paragraph_format.space_after = Pt(2)
    r_hdr = p_header.add_run("VIETNAM-NETHERLANDS MASTER’S PROGRAM IN DEVELOPMENT ECONOMICS (MDE)")
    r_hdr.bold = True
    r_hdr.font.size = Pt(13)
    r_hdr.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_td = doc.add_paragraph()
    p_td.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_td.paragraph_format.space_after = Pt(18)
    r_td = p_td.add_run("THESIS DESIGN")
    r_td.bold = True
    r_td.font.size = Pt(15)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(36)
    r_title = p_title.add_run("FACTORS AFFECTING CORPORATE CUSTOMERS’ DECISION TO CHOOSE BANK GUARANTEE SERVICES AT VIETNAM JOINT STOCK COMMERCIAL BANK FOR INDUSTRY AND TRADE (VIETINBANK)")
    r_title.bold = True
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_meta1 = doc.add_paragraph()
    p_meta1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta1.paragraph_format.space_after = Pt(4)
    r_sup = p_meta1.add_run("Supervisor(s): Dr. Hoang Thi Thuy Nga")
    r_sup.font.size = Pt(12)
    r_sup.bold = True

    p_meta2 = doc.add_paragraph()
    p_meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta2.paragraph_format.space_after = Pt(18)
    r_stu = p_meta2.add_run("Student: Dang Tu Linh, MDE Class 31")
    r_stu.font.size = Pt(12)
    r_stu.bold = True

    p_meta3 = doc.add_paragraph()
    p_meta3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_loc = p_meta3.add_run("Hanoi, 06/2026")
    r_loc.font.size = Pt(11)
    r_loc.italic = True

    doc.add_page_break()

    # Helper function for justified body paragraphs
    def add_body(text, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(space_after)
        p.add_run(text)
        return p

    def add_bullet(text, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(space_after)
        p.add_run("• " + text)
        return p

    def add_numbered(num_str, text, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(space_after)
        r_num = p.add_run(num_str + " ")
        r_num.bold = True
        p.add_run(text)
        return p

    # ==================== SECTION I: INTRODUCTION ====================
    doc.add_heading("I. Introduction", level=1)

    # 1.1 Research Rationales
    doc.add_heading("1.1. Research Rationales", level=2)

    add_body(
        "In modern commercial banking, bank guarantee operations represent a strategic off-balance sheet credit instrument designed to facilitate high-value commercial dealings, "
        "infrastructure projects, and international trade transactions. Under Article 18 of Decree No. 37/2015/ND-CP (amended and supplemented by Decree No. 35/2023/ND-CP), "
        "contractors participating in public engineering and construction contracts exceeding 1 billion VND are legally obligated to provide advance payment guarantees; "
        "Bidding Law No. 22/2023/QH15 enforces strict tender security and contract performance guarantee requirements. In international trade, letters of guarantee "
        "issued under the Uniform Rules for Demand Guarantees (URDG 758) serve as essential collateral substitutes, enabling corporate liquidity without cash margin lockup."
    )

    add_body(
        "From the corporate client's perspective, selecting a commercial bank to issue letters of guarantee is a multi-dimensional strategic decision. "
        "Enterprises evaluate not only fee pricing competitiveness, processing speed, and collateral margin flexibility, but also the issuing bank's financial reputation, "
        "legal advisory competence, relationship credit limits, and digital banking capabilities. A guarantee letter issued by a prestigious Big4 bank like VietinBank "
        "provides immediate credibility to project owners and foreign counterparties, ensuring seamless contract execution."
    )

    add_body(
        "From the commercial bank's perspective, bank guarantee issuance generates steady non-interest fee income while consuming lower regulatory capital under Basel credit conversion factors (CCF). "
        "However, commercial banks face intense interbank competition from domestic and foreign institutions. State Bank of Vietnam Circular No. 61/2024/TT-NHNN "
        "(effective April 1, 2025, replacing Circular No. 11/2022/TT-NHNN) establishes an updated legal framework for electronic guarantees (e-guarantees). "
        "Understanding what specific service attributes drive corporate customers' patronage decisions is vital for VietinBank to enhance product attractiveness and market share."
    )

    add_body(
        "VietinBank—one of Vietnam's Big4 state-owned commercial banks—maintains an extensive corporate banking network across 155 domestic branches. "
        "While previous empirical studies focus primarily on retail banking selection or general service quality, comprehensive quantitative models evaluating corporate guarantee selection decisions "
        "integrated with digital eFAST adoption remain scarce. Utilizing an established corporate survey dataset (n = 800) across 155 branches, this research investigates the primary drivers "
        "influencing corporate customers' patronage decisions. Hence, the author selects the topic: \"Factors Affecting Corporate Customers’ Decision to Choose Bank Guarantee Services at Vietnam Joint Stock Commercial Bank for Industry and Trade (VietinBank)\"."
    )

    # 1.2 General Objectives
    doc.add_heading("1.2. General Objectives", level=2)

    add_body(
        "The general objective of this study is to identify and assess the factors associated with corporate customers’ decision to choose VietinBank for bank guarantee services, "
        "and to propose managerial recommendations for improving VietinBank’s attractiveness and competitiveness in the corporate bank guarantee market."
    )

    # 1.3 Specific Objectives
    doc.add_heading("1.3. Specific Objectives", level=2)

    add_body("The study aims to fulfill the following four specific objectives:", space_after=4)
    
    sub_q = [
        ("1.", "Identify the key factors associated with corporate customers’ decision to choose VietinBank for bank guarantee services, based on relevant theories, previous empirical studies, and the characteristics of bank guarantee services."),
        ("2.", "Assess the direction and relative importance of these factors in explaining corporate customers’ selection decisions."),
        ("3.", "Examine whether corporate customers’ selection decisions differ across major firm characteristics, such as ownership type, firm size, operating experience, and types of bank guarantees used."),
        ("4.", "Propose managerial recommendations for VietinBank to improve its bank guarantee products and services and strengthen its ability to attract and retain corporate customers.")
    ]
    for num, q_text in sub_q:
        add_numbered(num, q_text, space_after=4)

    # 1.4 Thesis Structure
    doc.add_heading("1.4. Thesis Structure", level=2)

    add_body("This thesis comprises four main chapters:", space_after=4)
    add_bullet("Chapter 1: Introduction – Presents research rationales, background context, general and specific objectives, sub-questions, and thesis structure.")
    add_bullet("Chapter 2: Literature Review and Theoretical Framework – Synthesizes legal, economic, and banking selection foundations; reviews Financial Intermediation, Credit Risk Pricing, SERVQUAL, and Relationship Banking theories; analyzes empirical literature; and identifies key research gaps.")
    add_bullet("Chapter 3: Research Methodology and Empirical Design – Outlines the 7-step quantitative procedure (Descriptive, Cronbach's Alpha, EFA, Factor Scores, Pearson Correlation, VIF, OLS Regression, and Sub-group ANOVA/t-tests); defines dependent (DEC) and 7 independent variables; and describes the survey dataset (n = 800).")
    add_bullet("Chapter 4: Empirical Results, Discussion and Policy Recommendations – Reports EFA factor loadings, regression elasticities, sub-group difference tests, and formulates strategic managerial recommendations for VietinBank.", space_after=8)

    # ==================== SECTION II: LITERATURE REVIEW ====================
    doc.add_heading("II. Literature Review", level=1)

    # 2.1 Theoretical Foundations
    doc.add_heading("2.1. Theoretical Foundations and Legal Context", level=2)

    add_body(
        "Under Vietnamese law, State Bank of Vietnam Circular No. 61/2024/TT-NHNN (effective April 1, 2025, replacing Circular No. 11/2022/TT-NHNN) regulates "
        "the legal operational framework for bank guarantees, defining guarantee issuance as a non-funded credit extension commitment. "
        "Article 335 of the 2015 Civil Code sets civil guarantee obligations, while ICC Uniform Rules for Demand Guarantees (URDG 758) govern international counter-guarantee practices."
    )

    add_body("Corporate bank guarantee selection and patronage decisions are grounded in five foundational economic and management theories:", space_after=4)
    add_numbered("1.", "Financial Intermediation & Delegated Monitoring Theory (Diamond, 1984; Ramakrishnan & Thakor, 1984): Explains that commercial banks act as specialized information producers whose letters of guarantee signal corporate creditworthiness to third-party beneficiaries.")
    add_numbered("2.", "Credit Risk Pricing & Contingent Claim Theory (Merton, 1974; Stiglitz & Weiss, 1981): Establishes that bank guarantee fee competitiveness and margin requirements reflect underlying credit risk profiles and collateral protection.")
    add_numbered("3.", "SERVQUAL Service Quality Model (Parasuraman, Zeithaml & Berry, 1988): Framework measuring customer perception across service dimensions. Note that in corporate B2B banking, Bank Reputation (BANK_REP) is distinguished from operational Reliability.")
    add_numbered("4.", "Relationship Banking Theory (Boot, 2000; Berger & Udell, 1995): Demonstrates that multi-product credit ties and long-term history reduce asymmetric information, enabling customized guarantee limits and fee discounts.")
    add_numbered("5.", "Technology Acceptance Model & Digital Banking (Davis, 1989; Venkatesh et al., 2003): Explains how perceived usefulness and ease-of-use drive corporate adoption of digital channels (VietinBank eFAST e-guarantees).", space_after=8)

    # 2.2 Empirical Review & Gaps
    doc.add_heading("2.2. Empirical Review and Research Gaps", level=2)

    add_body(
        "Empirical literature spans corporate bank selection criteria (Turnbull & Gibbs, 1989; Narteh, 2013; Kaur et al., 2021; Zelie, 2023), "
        "Vietnamese B2B service quality (Phan Thi Hang Nga et al., 2024; Ho Dinh Phi et al., 2023; Nguyen et al., 2024), and trade finance practices (Al-Sabbagh & Al-Khathlan, 2018; Le Van Dung, 2021). "
        "Four key research gaps are identified: (i) limited empirical focus on off-balance sheet bank guarantees compared to general loan selection; "
        "(ii) lack of integrated models evaluating price competitiveness alongside digital eFAST adoption; "
        "(iii) absence of sub-group comparative analysis across corporate ownership types and guarantee product types in Vietnam; and "
        "(iv) lack of large-scale system-wide corporate empirical studies within VietinBank (n = 800)."
    )

    # ==================== SECTION III: EMPIRICAL ANALYSIS ====================
    doc.add_heading("III. Empirical Analysis", level=1)

    # 3.1 Methodology
    doc.add_heading("3.1. Methodology and Proposed Econometric Model", level=2)

    add_body(
        "Following the advisor's methodological guidance, the quantitative analysis follows a rigorous 7-step procedure: "
        "(1) Descriptive Statistics; (2) Reliability Analysis via Cronbach's Alpha (threshold >= 0.60, item-total correlation >= 0.30); "
        "(3) Exploratory Factor Analysis (EFA) with Varimax rotation (KMO >= 0.50, Eigenvalues >= 1.0, Cumulative Variance >= 50%); "
        "(4) Construction of Representative Factor Variables; (5) Pearson Correlation Matrix Analysis; "
        "(6) Multicollinearity Diagnostics via Variance Inflation Factor (VIF < 5); and "
        "(7) Multiple Linear Regression (OLS) and Sub-group Difference Testing (Independent Samples t-test and One-way ANOVA)."
    )

    add_body("The empirical regression equation is specified as follows:")

    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_after = Pt(10)
    r_eq = p_eq.add_run(
        "DEC = β0 + β1*COST_COMP + β2*PROC_SPEED + β3*DIGITAL_CONV + β4*BANK_REP + β5*RELATIONSHIP + β6*STAFF_QUAL + β7*COLL_POLICY + ε     (1)"
    )
    r_eq.bold = True
    r_eq.font.size = Pt(10.5)

    # Table 1: Variable Specifications
    p_t1_title = doc.add_paragraph()
    p_t1_title.paragraph_format.space_after = Pt(4)
    r_t1 = p_t1_title.add_run("Table 1: Variable Definitions and Quantitative Measurement Specifications (Option 2)")
    r_t1.bold = True
    r_t1.font.size = Pt(11)

    t1 = doc.add_table(rows=1, cols=5)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    col_w1 = [Inches(1.2), Inches(1.8), Inches(2.5), Inches(0.9), Inches(0.6)]

    hdr1 = t1.rows[0].cells
    hdr_titles1 = ["Code", "Variable Name", "Measurement Content / Data Source", "Type", "Sign"]
    for idx, t in enumerate(hdr_titles1):
        hdr1[idx].width = col_w1[idx]
        p = hdr1[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr1[idx], "003366")

    vars_info = [
        ("DEC", "Selection Decision", "Corporate preference and decision to select VietinBank as primary guarantee bank (Likert 1-5)", "Dependent (Y)", "N/A"),
        ("COST_COMP", "Price Competitiveness", "Fee reasonableness, competitive pricing, and discount policy (Likert 1-5)", "Independent (X1)", "+"),
        ("PROC_SPEED", "Processing Speed", "Underwriting turnaround time, fast issuance, and simple paperwork (Likert 1-5)", "Independent (X2)", "+"),
        ("DIGITAL_CONV", "Digital eFAST Convenience", "24/7 online submission, e-guarantee issuance, and digital tracking via eFAST (Likert 1-5)", "Independent (X3)", "+"),
        ("BANK_REP", "Bank Reputation", "Big4 brand prestige, financial capability, and universal market acceptance (Likert 1-5)", "Independent (X4)", "+"),
        ("RELATIONSHIP", "Relationship & Limits", "Credit limit flexibility, long-term relationship, and VIP customer care (Likert 1-5)", "Independent (X5)", "+"),
        ("STAFF_QUAL", "Staff Professionalism", "RM competency, legal advisory under Bidding Law & Circular 61 (Likert 1-5)", "Independent (X6)", "+"),
        ("COLL_POLICY", "Collateral & Margin Flexibility", "Flexible cash margin ratio, diverse pledged collateral types (Likert 1-5)", "Independent (X7)", "+")
    ]

    for code, name, desc, vtype, sign in vars_info:
        r_cells = t1.add_row().cells
        for i, w in enumerate(col_w1): r_cells[i].width = w
        
        p0 = r_cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p0.add_run(code).bold = True
        
        r_cells[1].paragraphs[0].add_run(name)
        r_cells[2].paragraphs[0].add_run(desc)
        
        p3 = r_cells[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(vtype)

        p4 = r_cells[4].paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.add_run(sign).bold = True

    # Table 1 Source Line (Gray Italic text below table)
    p_src = doc.add_paragraph()
    p_src.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_src.paragraph_format.space_before = Pt(4)
    p_src.paragraph_format.space_after = Pt(12)
    r_src = p_src.add_run("Source: Author's design based on VietinBank official corporate survey database (34 items, n = 800).")
    r_src.font.name = 'Times New Roman'
    r_src.font.size = Pt(9.5)
    r_src.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    r_src.italic = True

    # 3.2 Data Collection and Econometric Pre-Tests
    doc.add_heading("3.2. Data Collection and Econometric Pre-Tests", level=2)

    add_body(
        "The study utilizes an official survey dataset collected from n = 800 corporate customers across 155 VietinBank domestic branches nationwide. "
        "The survey instrument contains 34 Likert-scale items covering 7 independent factors and the dependent selection variable. "
        "The sample size (n = 800) substantially exceeds standard statistical guidelines (n >= 5 * total items = 170 for EFA; n >= 50 + 8*k = 106 for regression). "
        "Analytical steps executed in SPSS/Python include: Reliability analysis (Cronbach's Alpha >= 0.60), EFA (KMO >= 0.50, Varimax rotation), "
        "Factor Score extraction, Pearson correlation, VIF multicollinearity test (VIF < 5), Pooled OLS regression, and Sub-group ANOVA / t-tests across corporate ownership, size, age, and guarantee product types."
    )

    # 3.3 Findings and Interpretation
    doc.add_heading("3.3. Anticipated Findings and Interpretation", level=2)

    add_body("Based on theoretical frameworks and empirical literature, the following findings are anticipated:", space_after=4)
    add_bullet("All seven independent factors (COST_COMP, PROC_SPEED, DIGITAL_CONV, BANK_REP, RELATIONSHIP, STAFF_QUAL, COLL_POLICY) will exert positive, statistically significant impacts on corporate selection decisions (DEC).")
    add_bullet("Price Competitiveness (COST_COMP) and Bank Reputation (BANK_REP) are expected to emerge as the strongest drivers of corporate bank selection.")
    add_bullet("Digital eFAST Convenience (DIGITAL_CONV) will show a high positive coefficient, validating that digital e-guarantee capabilities enhance bank attractiveness.")
    add_bullet("Sub-group ANOVA tests will confirm significant differences in selection priority across corporate ownership types (SOEs vs Private vs FDI) and guarantee product types (Tender vs Performance vs Advance Payment Guarantees).", space_after=8)

    # ==================== SECTION IV: CONCLUSIONS ====================
    doc.add_heading("IV. Conclusions and Managerial Recommendations", level=1)

    add_body(
        "To ensure strict academic consistency, the managerial recommendations formulated in this section directly respond to the four specific research objectives (Q1 to Q4) set in Section 1.3:"
    )

    add_numbered(
        "1. Policy Response to Objective 1 (Factor Identification & Product Enhancement):",
        "VietinBank leadership should focus product enhancement strategies on the seven identified core drivers (COST_COMP, PROC_SPEED, DIGITAL_CONV, BANK_REP, RELATIONSHIP, STAFF_QUAL, COLL_POLICY), ensuring that guarantee products meet corporate operational needs."
    )

    add_numbered(
        "2. Policy Response to Objective 2 (Price Competitiveness & Service Optimization):",
        "Reflecting the strong positive impact of Price Competitiveness (COST_COMP) and Processing Speed (PROC_SPEED), VietinBank should implement transparent fee discounting schedules for high-volume clients and streamline internal credit approval workflows to reduce issuance turnaround time."
    )

    add_numbered(
        "3. Policy Response to Objective 3 (Sub-group Tailored Strategies):",
        "Addressing the empirical differences discovered across corporate ownership types, firm sizes, and guarantee product lines, VietinBank should design tailored corporate service packages: offering high credit limits for large SOEs, digital eFAST shortcuts for tech-savvy private SMEs, and specialized URDG 758 counter-guarantee structures for FDI enterprises."
    )

    add_numbered(
        "4. Policy Response to Objective 4 (Digital eFAST Transformation & Customer Retention):",
        "To strengthen VietinBank's competitive position and customer retention, executive management should launch a comprehensive Digital eFAST Promotion campaign, offering 24/7 e-guarantee issuing capabilities, integrated legal advisory support for Bidding Law compliance, and flexible collateral margin policies."
    )

    # ==================== SECTION V: REFERENCES ====================
    doc.add_heading("V. References", level=1)

    refs = [
        "Al-Sabbagh, M. and Al-Khathlan, K. (2018) 'Factors influencing corporate clients’ choice of commercial banks for trade finance services', Journal of Financial Services Marketing, 23(2), pp. 71–82.",
        "Baltagi, B.H. (2008) Econometric Analysis of Panel Data. 4th edn. Chichester: Wiley.",
        "Barru, D.J. (2005) 'How to Guarantee Contractor Performance on International Construction Projects: Comparing Surety Bonds with Bank Guarantees and Standby Letters of Credit', The George Washington International Law Review, 37(1), pp. 51–94.",
        "Berger, A.N. and Udell, G.F. (1995) 'Relationship Lending and Lines of Credit in Small Firm Finance', Journal of Business, 68(3), pp. 351–381.",
        "Bertrams, R.I.V.F. (2013) Bank Guarantees in International Trade. 4th edn. The Hague: Kluwer Law International.",
        "Boot, A.W.A. (2000) 'Relationship Banking: What Do We Know?', Journal of Financial Intermediation, 9(1), pp. 7–25.",
        "Carletti, E., Leonello, A. and Marquez, R. (2023) 'Loan guarantees, bank underwriting policies and financial fragility', Journal of Financial Economics, 149(2), pp. 260–295.",
        "Davis, F.D. (1989) 'Perceived Usefulness, Perceived Ease of Use, and User Acceptance of Information Technology', MIS Quarterly, 13(3), pp. 319–340.",
        "DeYoung, R. and Roland, K.P. (2001) 'Product Mix, Revenue Mix, and Risk at Commercial Banks', Journal of Financial Intermediation, 10(2), pp. 115–144.",
        "Diamond, D.W. (1984) 'Financial Intermediation and Delegated Monitoring', Review of Economic Studies, 51(3), pp. 393–414.",
        "Hassan, A.A. et al. (2018) 'The problems and abuse of performance bond in the construction industry', IOP Conference Series: Earth and Environmental Science, 143, p. 012045.",
        "Ho Dinh Phi et al. (2023) 'Effect of Service Quality on Customer Loyalty: the Mediation of Customer Satisfaction, and Corporate Reputation in Banking Industry', Eurasian Journal of Business and Management, 11(3), pp. 145–160.",
        "International Chamber of Commerce (2010) Uniform Rules for Demand Guarantees (URDG 758). ICC Publication No. 758. Paris: ICC.",
        "Kaur, M. et al. (2021) 'The determinants of bank selection criteria of SMEs: a fuzzy analytic hierarchy approach', Journal of Science and Technology Policy Management, 12(4), pp. 580–605.",
        "Le Van Dung (2021) 'The nature of payment guarantee relationships at credit institutions', Industry and Trade Magazine, 8(April), pp. 45–52.",
        "Merton, R.C. (1974) 'On the Pricing of Corporate Debt: The Risk Structure of Interest Rates', Journal of Finance, 29(2), pp. 449–470.",
        "Narteh, B. (2013) 'SME bank selection and patronage behaviour in the Ghanaian banking industry', Management Research Review, 36(11), pp. 1061–1080.",
        "Nguyen, H. et al. (2024) 'The impact of service innovation on customer satisfaction and customer loyalty: a case in Vietnamese retail banks', Future Business Journal, 10(1), p. 14.",
        "Nguyen Thi Nhung and Nguyen Duy Phu (2015) 'Payment guarantees at Vietnamese commercial banks', Development and Integration Magazine, 25(35), pp. 62–67.",
        "Oke, A.E. (2018) 'Bonding capability of Nigerian contracting firms', Engineering, Construction and Architectural Management, 25(8), pp. 1012–1024.",
        "Parasuraman, A., Zeithaml, V.A. and Berry, L.L. (1988) 'SERVQUAL: A Multiple-Item Scale for Measuring Consumer Perceptions of Service Quality', Journal of Retailing, 64(1), pp. 12–40.",
        "Phan Thi Hang Nga et al. (2024) 'Service quality, customer satisfaction and loyalty: a case study in Vietnamese SMEs', Cogent Business & Management, 11(1), p. 2304512.",
        "Ramakrishnan, R.T.S. and Thakor, A.V. (1984) 'Information Reliability and a Theory of Financial Intermediation', Review of Economic Studies, 51(3), pp. 415–432.",
        "State Bank of Vietnam (2024) Circular No. 61/2024/TT-NHNN dated December 31, 2024, providing regulations on bank guarantees (effective April 1, 2025). Hanoi: SBV.",
        "Stiglitz, J.E. and Weiss, A. (1981) 'Credit Rationing in Markets with Imperfect Information', American Economic Review, 71(3), pp. 393–410.",
        "Turnbull, P.W. and Gibbs, M.L. (1989) 'The Selection of Banks and Banking Services among Corporate Customers in South Africa', International Journal of Bank Marketing, 7(5), pp. 36–42.",
        "Venkatesh, V. et al. (2003) 'User Acceptance of Information Technology: Toward a Unified View', MIS Quarterly, 27(3), pp. 425–478.",
        "Zeithaml, V.A. (1988) 'Consumer Perceptions of Price, Quality, and Value: A Means-End Model and Synthesis of Evidence', Journal of Marketing, 52(3), pp. 2–22.",
        "Zeithaml, V.A., Berry, L.L. and Parasuraman, A. (1996) 'The Behavioral Consequences of Service Quality', Journal of Marketing, 60(2), pp. 31–46.",
        "Zelie, E.M. (2023) 'Factors determining bank selection by micro- and small-sized enterprises: evidence from Ethiopia', International Journal of Bank Marketing, 41(5), pp. 1120–1142."
    ]

    for ref in refs:
        p_ref = doc.add_paragraph()
        p_ref.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p_ref.paragraph_format.left_indent = Inches(0)
        p_ref.paragraph_format.first_line_indent = Inches(0)
        p_ref.paragraph_format.space_after = Pt(6)
        p_ref.add_run(ref).font.size = Pt(10.5)

    # Save output file (Overwrite workspace file DTL_Thesis_Design_NEU_MDE_Final.docx)
    output_filename = "c:/Users/nguyen.tuan.minh/Desktop/DTL-Master-Project/DTL_Thesis_Design_NEU_MDE_Final.docx"
    doc.save(output_filename)
    print(f"Successfully updated DTL_Thesis_Design_NEU_MDE_Final.docx with Advisor's review feedback at {output_filename}!")

if __name__ == "__main__":
    generate_neu_mde_final_docx()
