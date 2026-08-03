import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def set_cell_padding(cell, top=120, bottom=120, left=150, right=150):
    tcPr = cell._element.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for side, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{side}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_bottom_border_to_paragraph(paragraph, color_hex="003366", size="8"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def create_survey_docx():
    doc = docx.Document()

    # 1. Page Margins (Standard 1.0 inch on all sides)
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # 2. Base Normal Style
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(4)

    # ==================== INSTITUTIONAL HEADER ====================
    p_hdr = doc.add_paragraph()
    p_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_hdr.paragraph_format.space_after = Pt(2)
    r_hdr1 = p_hdr.add_run("NGÂN HÀNG TMCP CÔNG THƯƠNG VIỆT NAM (VIETINBANK)")
    r_hdr1.bold = True
    r_hdr1.font.size = Pt(12)
    r_hdr1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_proj = doc.add_paragraph()
    p_proj.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_proj.paragraph_format.space_after = Pt(10)
    r_hdr2 = p_proj.add_run("DỰ ÁN KHẢO SÁT Ý KIẾN KHÁCH HÀNG DOANH NGHIỆP VỀ DỊCH VỤ BẢO LÃNH")
    r_hdr2.bold = True
    r_hdr2.font.size = Pt(11)
    r_hdr2.font.color.rgb = RGBColor(0x00, 0x55, 0x99)
    add_bottom_border_to_paragraph(p_proj, color_hex="003366", size="8")

    doc.add_paragraph()

    # Document Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_after = Pt(4)
    r_title = p_title.add_run("PHIẾU KHẢO SÁT CHÍNH THỨC")
    r_title.bold = True
    r_title.font.size = Pt(15)
    r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("CÁC YẾU TỐ ẢNH HƯỞNG ĐẾN QUYẾT ĐỊNH LỰA CHỌN SẢN PHẨM BẢO LÃNH NGÂN HÀNG CỦA CÁC DOANH NGHIỆP TẠI VIETINBANK")
    r_sub.bold = True
    r_sub.font.size = Pt(11.5)
    r_sub.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # Intro Letter Box (Separate paragraphs to prevent stretched JUSTIFY text!)
    p_i1 = doc.add_paragraph()
    p_i1.paragraph_format.space_after = Pt(4)
    r_i1 = p_i1.add_run("Kính gửi: Quý Doanh nghiệp,")
    r_i1.bold = True
    r_i1.italic = True

    p_i2 = doc.add_paragraph()
    p_i2.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_i2.paragraph_format.line_spacing = 1.15
    p_i2.paragraph_format.space_after = Pt(4)
    r_i2 = p_i2.add_run(
        "Nhằm nâng cao chất lượng dịch vụ, tối ưu hóa quy trình và thiết kế các gói sản phẩm bảo lãnh đáp ứng tốt nhất nhu cầu của Quý Doanh nghiệp, "
        "VietinBank tiến hành cuộc khảo sát này. Kính mong Quý Doanh nghiệp dành chút thời gian quý báu để đưa ra những đánh giá khách quan."
    )
    r_i2.italic = True

    p_i3 = doc.add_paragraph()
    p_i3.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p_i3.paragraph_format.line_spacing = 1.15
    p_i3.paragraph_format.space_after = Pt(4)
    r_i3 = p_i3.add_run("VietinBank cam kết toàn bộ thông tin do Quý Doanh nghiệp cung cấp sẽ được giữ bí mật tuyệt đối và chỉ sử dụng cho mục đích phân tích tổng hợp.")
    r_i3.italic = True

    p_i4 = doc.add_paragraph()
    p_i4.paragraph_format.space_after = Pt(14)
    r_i4 = p_i4.add_run("Xin chân thành cảm ơn sự hợp tác của Quý Doanh nghiệp!")
    r_i4.italic = True
    r_i4.bold = True

    # ==================== PART I: DEMOGRAPHICS ====================
    h_part1 = doc.add_paragraph()
    r_p1 = h_part1.add_run("PHẦN I: THÔNG TIN CHUNG VỀ DOANH NGHIỆP")
    r_p1.bold = True
    r_p1.font.size = Pt(12)
    r_p1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    demo_items = [
        ("1. Loại hình Doanh nghiệp:", [
            "[  ] DNTN / Công ty TNHH",
            "[  ] Công ty Cổ phần",
            "[  ] Doanh nghiệp Nhà nước / FDI",
            "[  ] Loại hình khác"
        ]),
        ("2. Quy mô Doanh thu trung bình hằng năm:", [
            "[  ] Dưới 20 tỷ VNĐ",
            "[  ] Từ 20 đến 100 tỷ VNĐ",
            "[  ] Từ 100 đến 500 tỷ VNĐ",
            "[  ] Trên 500 tỷ VNĐ"
        ]),
        ("3. Thâm niên hoạt động của Doanh nghiệp:", [
            "[  ] Dưới 3 năm",
            "[  ] Từ 3 đến 5 năm",
            "[  ] Từ 5 đến 10 năm",
            "[  ] Trên 10 năm"
        ]),
        ("4. Loại hình Bảo lãnh thường xuyên sử dụng tại VietinBank (Có thể chọn nhiều mục):", [
            "[  ] Bảo lãnh dự thầu (TG)",
            "[  ] Bảo lãnh thực hiện hợp đồng (PG)",
            "[  ] Bảo lãnh tạm ứng / thanh toán (BG)",
            "[  ] Bảo lãnh bảo hành / Tái bảo lãnh (RG)"
        ]),
        ("5. Chức danh / Vị trí của Người đại diện trả lời phiếu:", [
            "[  ] Ban Giám đốc / CFO",
            "[  ] Kế toán trưởng / Trưởng phòng Tài chính",
            "[  ] Trưởng phòng Thầu / Mua hàng",
            "[  ] Chuyên viên phụ trách bảo lãnh"
        ])
    ]

    for q_title, opts in demo_items:
        p_q = doc.add_paragraph()
        p_q.paragraph_format.space_after = Pt(2)
        r_q = p_q.add_run(q_title)
        r_q.bold = True

        p_opts = doc.add_paragraph()
        p_opts.paragraph_format.left_indent = Inches(0.3)
        p_opts.paragraph_format.space_after = Pt(6)
        p_opts.add_run("   ".join(opts))

    doc.add_paragraph()

    # ==================== PART II: LIKERT SCALE TABLE ====================
    h_part2 = doc.add_paragraph()
    r_p2 = h_part2.add_run("PHẦN II: NỘI DUNG ĐÁNH GIÁ (THANG ĐO LIKERT 5 MỨC ĐỘ)")
    r_p2.bold = True
    r_p2.font.size = Pt(12)
    r_p2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_guide = doc.add_paragraph()
    p_guide.paragraph_format.space_after = Pt(10)
    r_gd = p_guide.add_run(
        "Quy ước mức độ đánh giá: "
        "1 – Hoàn toàn không đồng ý  |  2 – Không đồng ý  |  3 – Phân vân / Trung lập  |  4 – Đồng ý  |  5 – Hoàn toàn đồng ý"
    )
    r_gd.font.size = Pt(10)
    r_gd.italic = True

    # Main Likert Table
    table = doc.add_table(rows=1, cols=7)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    col_widths = [Inches(0.8), Inches(4.1), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35), Inches(0.35)]

    # Header Row
    hdr_cells = table.rows[0].cells
    hdr_titles = ["Mã biến", "Nội dung phát biểu đánh giá", "1", "2", "3", "4", "5"]
    for idx, title in enumerate(hdr_titles):
        hdr_cells[idx].width = col_widths[idx]
        set_cell_padding(hdr_cells[idx], top=140, bottom=140, left=100, right=100)
        p = hdr_cells[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(title)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr_cells[idx], "003366")

    # Data Structure: 10 Independent + 1 Dependent = 34 Items Total
    sections = [
        ("A. CÁC YẾU TỐ ĐỘC LẬP (CÁC TIÊU CHÍ CHẤT LƯỢNG DỊCH VỤ)", [
            ("I. Chi phí & Biểu phí bảo lãnh (COST)", [
                ("COST1", "Mức phí phát hành bảo lãnh tại VietinBank là hợp lý so với chất lượng dịch vụ nhận được."),
                ("COST2", "Biểu phí bảo lãnh của VietinBank có tính cạnh tranh cao so với các ngân hàng thương mại khác."),
                ("COST3", "VietinBank có chính sách ưu đãi và giảm phí bảo lãnh linh hoạt cho các khách hàng thường xuyên.")
            ]),
            ("II. Yêu cầu Tài sản đảm bảo & Tỷ lệ Ký quỹ (COL)", [
                ("COL1", "VietinBank áp dụng tỷ lệ ký quỹ bảo lãnh linh hoạt (có chính sách giảm hoặc miễn ký quỹ)."),
                ("COL2", "VietinBank chấp nhận đa dạng các loại tài sản đảm bảo (BĐS, máy móc, dòng tiền hợp đồng...)."),
                ("COL3", "Thủ tục định giá và thế chấp tài sản đảm bảo tại VietinBank được thực hiện nhanh chóng, thuận tiện.")
            ]),
            ("III. Thời gian & Tốc độ xử lý hồ sơ (SPE)", [
                ("SPE1", "Thời gian thẩm định và phê duyệt hạn mức bảo lãnh tại VietinBank rất nhanh chóng."),
                ("SPE2", "Quy trình thủ tục hồ sơ đề nghị cấp bảo lãnh tại VietinBank được đơn giản hóa tối đa."),
                ("SPE3", "Thời gian từ khi nộp đủ hồ sơ đến khi nhận Thư bảo lãnh gốc đáp ứng kịp thời tiến độ hợp đồng.")
            ]),
            ("IV. Uy tín & Thương hiệu VietinBank (REP)", [
                ("REP1", "VietinBank là ngân hàng thương mại lớn, có uy tín và thương hiệu mạnh hàng đầu Việt Nam."),
                ("REP2", "Thư bảo lãnh do VietinBank phát hành luôn được các Chủ đầu tư và Bên mời thầu tin tưởng chấp nhận."),
                ("REP3", "Năng lực tài chính vững mạnh của VietinBank giúp gia tăng uy tín cho doanh nghiệp khi giao dịch.")
            ]),
            ("V. Trình độ & Thái độ phục vụ của Cán bộ (STA)", [
                ("STA1", "Cán bộ quan hệ khách hàng (RM) VietinBank có năng lực chuyên môn sâu về nghiệp vụ bảo lãnh."),
                ("STA2", "Cán bộ VietinBank am hiểu đặc thù ngành nghề kinh doanh và luôn tư vấn giải pháp phù hợp."),
                ("STA3", "Thái độ phục vụ của nhân viên VietinBank luôn tận tình, chuyên nghiệp và hỗ trợ kịp thời.")
            ]),
            ("VI. Mối quan hệ tín dụng sẵn có (REL)", [
                ("REL1", "Doanh nghiệp có lịch sử tín dụng tốt và mối quan hệ gắn kết lâu năm với VietinBank."),
                ("REL2", "Việc đang mở tài khoản thanh toán / trả lương tại VietinBank giúp thủ tục bảo lãnh thuận lợi hơn."),
                ("REL3", "Doanh nghiệp dễ dàng được cấp hạn mức bảo lãnh nhờ đã từng sử dụng các dịch vụ tài trợ khác.")
            ]),
            ("VII. Mức độ Chuyển đổi số & Bảo lãnh điện tử (DIG - Biến mới 1)", [
                ("DIG1", "Doanh nghiệp có thể dễ dàng nộp đề nghị cấp bảo lãnh online 24/7 qua VietinBank eFAST."),
                ("DIG2", "VietinBank phát hành Thư bảo lãnh điện tử ứng dụng chữ ký số CA nhanh chóng, không cần bản giấy."),
                ("DIG3", "Tính năng tra cứu mã QR xác thực Thư bảo lãnh thật/giả online trên hệ thống rất tiện ích và an toàn.")
            ]),
            ("VIII. Thiết kế Giải pháp Bảo lãnh theo yêu cầu (CUS - Biến mới 2)", [
                ("CUS1", "VietinBank có khả năng thiết kế linh hoạt mẫu thư bảo lãnh theo yêu cầu riêng của Chủ đầu tư."),
                ("CUS2", "VietinBank cung cấp cấu trúc bảo lãnh đa dạng (bảo lãnh tạm ứng, bảo lãnh thanh toán, đối ứng...)."),
                ("CUS3", "Ngân hàng sẵn sàng điều chỉnh các điều khoản bảo lãnh phù hợp với tiến độ thi công thực tế.")
            ]),
            ("IX. Năng lực Tư vấn Pháp lý & Quản trị Rủi ro (RSK - Biến mới 3)", [
                ("RSK1", "VietinBank hỗ trợ tư vấn rà soát tính pháp lý của hợp đồng gốc trước khi phát hành bảo lãnh."),
                ("RSK2", "Cán bộ VietinBank tư vấn cấu trúc điều khoản bảo lãnh giúp phòng ngừa rủi ro bị tịch thu vô lý."),
                ("RSK3", "VietinBank áp dụng tốt các tập quán bảo lãnh quốc tế (URDG 758) giúp bảo vệ quyền lợi doanh nghiệp.")
            ]),
            ("X. Mạng lưới Ngân hàng Đại lý Quốc tế (NET - Biến mới 4)", [
                ("NET1", "VietinBank có mạng lưới ngân hàng đại lý rộng khắp toàn cầu, hỗ trợ tốt cho giao dịch quốc tế."),
                ("NET2", "Thư bảo lãnh đối ứng (Counter-Guarantee) của VietinBank dễ dàng được các đối tác nước ngoài chấp nhận."),
                ("NET3", "VietinBank hỗ trợ hiệu quả các thủ tục bảo lãnh bằng ngoại tệ đối với các hợp đồng xuất nhập khẩu.")
            ])
        ]),
        ("B. QUYẾT ĐỊNH LỰA CHỌN VIETINBANK (BIẾN MỤC TIÊU Y)", [
            ("XI. Quyết định Lựa chọn & Sự hài lòng của Doanh nghiệp (DEC)", [
                ("DEC1", "Doanh nghiệp của bạn luôn ưu tiên lựa chọn VietinBank để phát hành các thư bảo lãnh."),
                ("DEC2", "Doanh nghiệp sẵn sàng tăng quy mô và tần suất sử dụng dịch vụ bảo lãnh tại VietinBank trong tương lai."),
                ("DEC3", "Doanh nghiệp sẵn sàng giới thiệu dịch vụ bảo lãnh của VietinBank cho các đối tác, khách hàng khác."),
                ("DEC4", "Nhìn chung, VietinBank là lựa chọn tối ưu nhất của doanh nghiệp khi phát sinh nhu cầu bảo lãnh.")
            ])
        ])
    ]

    for part_title, group_list in sections:
        # Part Banner Row - MERGE ACROSS ALL 7 COLUMNS!
        r_part = table.add_row().cells
        cell_part = r_part[0].merge(r_part[6])
        set_cell_padding(cell_part, top=100, bottom=100, left=120, right=120)
        p_part = cell_part.paragraphs[0]
        r_pt = p_part.add_run(part_title)
        r_pt.bold = True
        r_pt.font.size = Pt(10.5)
        r_pt.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
        set_cell_background(cell_part, "E6F0FA")

        for grp_name, item_list in group_list:
            # Group Subheader Row - MERGE ACROSS ALL 7 COLUMNS!
            r_grp = table.add_row().cells
            cell_grp = r_grp[0].merge(r_grp[6])
            set_cell_padding(cell_grp, top=80, bottom=80, left=120, right=120)
            p_grp = cell_grp.paragraphs[0]
            r_gt = p_grp.add_run(grp_name)
            r_gt.bold = True
            r_gt.font.size = Pt(10)
            set_cell_background(cell_grp, "F4F4F4")

            for code, statement in item_list:
                r_item = table.add_row().cells
                for i, w in enumerate(col_widths): r_item[i].width = w
                
                # Apply cell padding
                for c_cell in r_item:
                    set_cell_padding(c_cell, top=80, bottom=80, left=100, right=100)

                p_code = r_item[0].paragraphs[0]
                p_code.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p_code.add_run(code).bold = True

                p_stmt = r_item[1].paragraphs[0]
                p_stmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
                p_stmt.paragraph_format.line_spacing = 1.15
                p_stmt.add_run(statement)

                for c_idx in range(2, 7):
                    p_box = r_item[c_idx].paragraphs[0]
                    p_box.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    p_box.add_run("(  )")

    # Thank you footer
    doc.add_paragraph()
    p_foot = doc.add_paragraph()
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(12)
    r_ft = p_foot.add_run("XIN CHÂN THÀNH CẢM ƠN QUÝ DOANH NGHIỆP ĐÃ HOÀN THÀNH PHIẾU KHẢO SÁT!")
    r_ft.bold = True
    r_ft.font.size = Pt(11)
    r_ft.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    # Save output DOCX
    out_path = "c:/Users/nguyen.tuan.minh/Desktop/DTL-Master-Project/Phieu_Khao_Sat_Chinh_Thuc_VietinBank.docx"
    doc.save(out_path)
    print(f"Successfully fixed cell merging and stretched text in survey DOCX at {out_path}")

if __name__ == "__main__":
    create_survey_docx()
