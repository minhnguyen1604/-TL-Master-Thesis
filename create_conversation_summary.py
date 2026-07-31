import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def create_summary_docx():
    doc = docx.Document()
    
    # Margins
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    # Styles
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(11.5)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)

    # Title
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("TỔNG HỢP TOÀN BỘ NỘI DUNG TRAO ĐỔI & ĐỊNH HƯỚNG LUẬN VĂN THẠC SĨ\n")
    r_title.bold = True
    r_title.font.size = Pt(16)
    r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("Chủ đề: Nghiên cứu Dịch vụ Bảo lãnh Ngân hàng dành cho Khách hàng Doanh nghiệp tại VietinBank\n")
    r_sub.italic = True
    r_sub.bold = True
    r_sub.font.size = Pt(13)
    r_sub.font.color.rgb = RGBColor(0x00, 0x55, 0x99)

    doc.add_paragraph()

    # SECTION 1: TÊN ĐỀ TÀI
    h1 = doc.add_paragraph()
    r = h1.add_run("I. TÊN ĐỀ TÀI & PHẠM VI NGHIÊN CỨU TỔNG THỂ")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.add_run("• Tên đề tài Tiếng Việt chính thức:\n").bold = True
    p.add_run("  \"Các yếu tố ảnh hưởng đến phát triển hoạt động bảo lãnh ngân hàng cho khách hàng doanh nghiệp tại Ngân hàng Thương mại Cổ phần Công Thương Việt Nam\"\n\n")
    p.add_run("• Tên đề tài Tiếng Anh tương ứng:\n").bold = True
    p.add_run("  \"Factors Affecting the Development of Corporate Bank Guarantee Services at Vietnam Joint Stock Commercial Bank for Industry and Trade\"\n\n")
    p.add_run("• Tầm vóc đề tài:\n").bold = True
    p.add_run("  Đề tài mang tính bao quát chiến lược toàn hệ thống VietinBank, kết hợp giữa quy mô doanh số phát hành, hiệu quả thu nhập phí bảo lãnh, và định hướng chuyển đổi số bảo lãnh điện tử (eFAST).\n")

    # SECTION 2: PHƯƠNG ÁN 1
    h2 = doc.add_paragraph()
    r = h2.add_run("II. PHƯƠNG ÁN 1: NGHIÊN CỨU BẰNG DỮ LIỆU KHẢO SÁT (PRIMARY SURVEY DATA)")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.add_run("1. Bản chất phương pháp:\n").bold = True
    p.add_run("  Phát phiếu khảo sát dạng thang đo Likert 5 điểm (1: Hoàn toàn không đồng ý -> 5: Hoàn toàn đồng ý) cho người đại diện doanh nghiệp (CFO, Kế toán trưởng, Trưởng phòng thầu).\n\n")
    p.add_run("2. Cấu trúc Biến mục tiêu Y (Quyết định lựa chọn - DEC):\n").bold = True
    p.add_run("  Gồm 4 biến quan sát đo lường 4 giai đoạn lòng trung thành theo khung lý thuyết Oliver (1999) & Zeithaml et al. (1996):\n")
    p.add_run("  - DEC1: Doanh nghiệp luôn ưu tiên lựa chọn VietinBank để phát hành thư bảo lãnh (Trung thành Nhận thức).\n")
    p.add_run("  - DEC2: Doanh nghiệp sẵn sàng tăng quy mô và tần suất sử dụng bảo lãnh VietinBank trong tương lai (Trung thành Ý định).\n")
    p.add_run("  - DEC3: Doanh nghiệp sẵn sàng giới thiệu dịch vụ bảo lãnh VietinBank cho các đối tác khác (Trung thành Lan tỏa).\n")
    p.add_run("  - DEC4: Nhìn chung, VietinBank là lựa chọn tối ưu nhất khi phát sinh nhu cầu bảo lãnh (Trung thành Thái độ so với đối thủ).\n\n")
    p.add_run("3. Các biến độc lập (X1 -> X7):\n").bold = True
    p.add_run("  Bao gồm COST (Chi phí), COL (Tài sản đảm bảo), SPE (Tốc độ xử lý), REP (Uy tín), STA (Cán bộ RM), REL (Mối quan hệ), và DIG (Chuyển đổi số & e-Guarantee - Biến mới).\n")

    # SECTION 3: PHƯƠNG ÁN 2
    h3 = doc.add_paragraph()
    r = h3.add_run("III. PHƯƠNG ÁN 2: NGHIÊN CỨU BẰNG DỮ LIỆU GIAO DỊCH THỰC TẾ (SECONDARY TRANSACTION DATA)")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.add_run("1. Nguồn dữ liệu thực tế:\n").bold = True
    p.add_run("  Trích xuất từ 2 File Excel dữ liệu gốc của VietinBank:\n")
    p.add_run("  - File 1: doanh_số_BL_6T2026_theo_ĐCTC_tại_CN_108_ (2).xlsx (Chi tiết Doanh số cấp bảo lãnh quy đổi VNĐ).\n")
    p.add_run("  - File 2: Phí BL chi tiết từng CIF tại CN PK 98.xlsx (Chi tiết Thu phí bảo lãnh, Loại bảo lãnh, Chi nhánh, Khu vực).\n")
    p.add_run("  -> Ghép dữ liệu từ 2 file về 1 Bảng duy nhất theo Khóa chung là Mã Số CIF.\n\n")

    p.add_run("2. Các phương án xác định Biến mục tiêu (Y):\n").bold = True
    p.add_run("  - Phương án 1 (Biến đơn): Y = LN(Thu_Phi_BL_VND)\n")
    p.add_run("  - Phương án 2 (Kỹ thuật PCA): Gộp Doanh số và Thu phí bằng Chuẩn hóa Z-score & Phân tích nhân tố PCA -> Y_PERFORMANCE\n")
    p.add_run("  - Phương án 3 (Tỷ lệ Thu phí): Y = (Số tiền Phí thu được VNĐ / Doanh số Bảo lãnh cấp ra VNĐ) * 100%\n")
    p.add_run("  - Phương án 4 (Tổng Logarit): Y = LN(Doanh số Bảo lãnh cấp ra VNĐ) + LN(Số tiền Phí bảo lãnh thu được VNĐ)\n\n")

    p.add_run("3. Danh mục 6 Biến độc lập (X1 -> X6) chuẩn hóa trong mô hình:\n").bold = True
    p.add_run("  - TY_LE_PHI (X1): Tỷ lệ phí bảo lãnh thực tế (%)\n")
    p.add_run("  - TAN_SUAT_GD (X2): Tổng số lượng thư bảo lãnh/hợp đồng thực hiện trong kỳ\n")
    p.add_run("  - LOAI_BL (X3): Phân loại dòng bảo lãnh (BG, PG, TG, RG)\n")
    p.add_run("  - PHAN_KHUC (X4): Phân khúc khách hàng (98 - Định chế tài chính, SMEs, KHDN Lớn)\n")
    p.add_run("  - KHU_VUC (X5): Địa bàn vùng miền của Chi nhánh (KV1, KV2, KV3, KV7...)\n")
    p.add_run("  - KENH_EFAST (X6): Biến giả Chuyển đổi số (1 = giao dịch qua eFAST, 0 = làm tại quầy)\n")

    # SECTION 4: HÀM Ý QUẢN TRỊ
    h4 = doc.add_paragraph()
    r = h4.add_run("IV. ĐÓNG GÓP HÀM Ý QUẢN TRỊ & GIẢI PHÁP CHO VIETINBANK")
    r.bold = True
    r.font.size = Pt(13)
    r.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p = doc.add_paragraph()
    p.paragraph_format.line_spacing = 1.2
    p.add_run("1. Tối ưu hóa Biểu phí & Phân hóa giá:\n").bold = True
    p.add_run("  Phân tích độ co giãn của phí để tìm mức phí tối ưu cho từng phân khúc; giảm phí cho khách hàng xếp hạng tín nhiệm cao và khách hàng phát hành qua eFAST.\n\n")
    p.add_run("2. Đẩy mạnh Số hóa & Phê duyệt tự động STP:\n").bold = True
    p.add_run("  Đề xuất nâng hạn mức duyệt tự động trên VietinBank eFAST cho các hợp đồng dưới 5 tỷ VNĐ không cần qua thẩm định thủ công.\n\n")
    p.add_run("3. Cơ chế Ký quỹ linh hoạt:\n").bold = True
    p.add_run("  Giảm hoặc miễn 100% ký quỹ cho bảo lãnh dự thầu (TG) với doanh nghiệp uy tín để giải phóng dòng tiền cho nhà thầu.\n")

    # Save
    out_path = "c:/Users/nguyen.tuan.minh/Desktop/DTL-Master-Project/Tong_Hop_Bao_Cao_De_Tai_Master_VietinBank.docx"
    doc.save(out_path)
    print(f"Saved DOCX summary at {out_path}")

if __name__ == "__main__":
    create_summary_docx()
