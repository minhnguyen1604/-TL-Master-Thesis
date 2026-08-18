import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_background(cell, fill_hex):
    tcPr = cell._element.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_hex)
    tcPr.append(shd)

def add_bottom_border_to_paragraph(paragraph, color_hex="888888", size="6"):
    pPr = paragraph._element.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), size)
    bottom.set(qn('w:space'), '4')
    bottom.set(qn('w:color'), color_hex)
    pBdr.append(bottom)
    pPr.append(pBdr)

def generate_report_docx():
    doc = docx.Document()
    
    # 1. Page Margins (1.0 inch = 2.54 cm)
    section = doc.sections[0]
    section.top_margin = Inches(1.0)
    section.bottom_margin = Inches(1.0)
    section.left_margin = Inches(1.0)
    section.right_margin = Inches(1.0)

    # 2. Running Header Configuration (Different First Page)
    section.different_first_page_header_footer = True
    header = section.header
    p_head = header.paragraphs[0]
    p_head.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    r_head = p_head.add_run("Báo cáo Tiếp thu & Giải trình Góp ý Giảng viên Hướng dẫn – Đặng Tú Linh MDE31")
    r_head.font.name = 'Times New Roman'
    r_head.font.size = Pt(9.5)
    r_head.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    r_head.italic = True
    add_bottom_border_to_paragraph(p_head, color_hex="888888", size="6")

    # Configure Styles
    styles = doc.styles
    style_normal = styles['Normal']
    style_normal.font.name = 'Times New Roman'
    style_normal.font.size = Pt(12)
    style_normal.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    style_normal.paragraph_format.line_spacing = 1.15
    style_normal.paragraph_format.space_after = Pt(6)

    style_h1 = styles['Heading 1']
    style_h1.font.name = 'Times New Roman'
    style_h1.font.size = Pt(14)
    style_h1.font.bold = True
    style_h1.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    style_h1.paragraph_format.space_before = Pt(14)
    style_h1.paragraph_format.space_after = Pt(6)

    style_h2 = styles['Heading 2']
    style_h2.font.name = 'Times New Roman'
    style_h2.font.size = Pt(12.5)
    style_h2.font.bold = True
    style_h2.font.color.rgb = RGBColor(0x00, 0x33, 0x66)
    style_h2.paragraph_format.space_before = Pt(10)
    style_h2.paragraph_format.space_after = Pt(4)

    # Header / Title Block
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_inst = p_inst.add_run("CHƯƠNG TRÌNH CAO HỌC KINH TẾ PHÁT TRIỂN VIỆT NAM - HÀ LAN (MDE)\nTRƯỜNG ĐẠI HỌC KINH TẾ QUỐC DÂN")
    r_inst.bold = True
    r_inst.font.size = Pt(12)
    r_inst.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    r_title = p_title.add_run("BÁO CÁO TIẾP THU, GIẢI TRÌNH VÀ TỔNG HỢP NỘI DUNG CHỈNH SỬA\nTHEO GÓP Ý CỦA GIẢNG VIÊN HƯỚNG DẪN")
    r_title.bold = True
    r_title.font.size = Pt(14)
    r_title.font.color.rgb = RGBColor(0x00, 0x33, 0x66)

    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_meta.paragraph_format.space_after = Pt(18)
    r_meta = p_meta.add_run(
        "Đề tài: Factors Affecting Corporate Customers’ Decision to Choose Bank Guarantee Services at VietinBank\n"
        "Học viên: Đặng Tú Linh – Lớp MDE Khóa 31\n"
        "Giảng viên hướng dẫn: TS. Hoàng Thị Thúy Nga"
    )
    r_meta.italic = True
    r_meta.font.size = Pt(11)

    def add_body(text, space_after=6):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(space_after)
        p.add_run(text)
        return p

    def add_bullet(text, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(space_after)
        p.add_run("• " + text)
        return p

    def add_numbered(num_str, text, space_after=4):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.space_after = Pt(space_after)
        r_num = p.add_run(num_str + " ")
        r_num.bold = True
        p.add_run(text)
        return p

    # ==================== PHẦN I ====================
    doc.add_heading("PHẦN I: TỔNG HỢP VÀ ĐÁNH GIÁ Ý KIẾN CỦA GIẢNG VIÊN HƯỚNG DẪN", level=1)

    doc.add_heading("1.1. Các nội dung Giảng viên hướng dẫn ĐỒNG TÌNH (Phê duyệt)", level=2)
    add_body(
        "Giảng viên hướng dẫn đánh giá cao định hướng nghiên cứu tổng thể và hoàn toàn đồng thuận với các trụ cột phương pháp luận sau:"
    )
    add_bullet("Về Phương pháp nghiên cứu: Đồng ý sử dụng mô hình Hồi quy tuyến tính đa biến (Multiple Linear Regression – OLS) để ước lượng tác động của các nhân tố đến quyết định lựa chọn dịch vụ bảo lãnh.")
    add_bullet("Về Quy trình phân tích định lượng 7 bước: Phê duyệt 100% quy trình tuần tự: Thống kê mô tả → Đánh giá độ tin cậy Cronbach’s Alpha → Khám phá nhân tố EFA → Trích xuất biến đại diện (Factor Scores) → Phân tích tương quan Pearson → Kiểm tra đa cộng tuyến (VIF) → Hồi quy OLS và kiểm định giả thuyết nghiên cứu.")
    add_bullet("Về Bối cảnh và Nguồn dữ liệu: Nhất trí với việc khai thác bộ dữ liệu khảo sát thực tế quy mô n = 800 khách hàng doanh nghiệp trên 155 chi nhánh toàn quốc của Ngân hàng TMCP Công thương Việt Nam (VietinBank).")

    doc.add_heading("1.2. Các nội dung Giảng viên hướng dẫn YÊU CẦU ĐIỀU CHỈNH & Hướng xử lý", level=2)
    add_body("Học viên đã nghiêm túc tiếp thu 5 nội dung điều chỉnh trọng tâm của Giảng viên hướng dẫn và triển khai xử lý như sau:")

    # Table 1: Feedback Analysis Table
    t1 = doc.add_table(rows=1, cols=3)
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.autofit = False
    col_w1 = [Inches(1.8), Inches(2.2), Inches(2.5)]

    hdr1 = t1.rows[0].cells
    hdr_titles1 = ["Ý kiến Góp ý của Giảng viên", "Phân tích Học thuật & Nguyên nhân", "Phương án Tiếp thu & Xử lý"]
    for idx, t in enumerate(hdr_titles1):
        hdr1[idx].width = col_w1[idx]
        p = hdr1[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr1[idx], "003366")

    fb_items = [
        (
            "1. Rút gọn số lượng biến độc lập:\nNên chọn khoảng 6–7 yếu tố đặc trưng, tránh dàn trải. Lưu ý Bank Reputation không phải Reliability trong SERVQUAL.",
            "Bản draft cũ có 10 biến làm mô hình bị phân tán, giảm bậc tự do và dễ dính đa cộng tuyến.",
            "Gom nhóm và rút gọn xuống đúng 7 biến độc lập trọng tâm (nằm đúng khung 6-7 biến), giữ lại biến Chuyển đổi số eFAST và tách riêng Uy tín thương hiệu (BANK_REP)."
        ),
        (
            "2. Điều chỉnh biến COST & Chiều tác động (+):\nThang đo đo mức độ hợp lý/cạnh tranh của phí nhưng giả thuyết lại để dấu âm (-). Cần đổi tên thành Price Competitiveness và kỳ vọng tác động dương (+).",
            "Điểm Likert 5 thể hiện 'Phí rất hợp lý và cạnh tranh'. Điểm càng cao thì càng hấp dẫn doanh nghiệp chọn VietinBank nên dấu kỳ vọng bắt buộc phải là DƯƠNG (+).",
            "Đổi tên biến thành COST_COMP (Price Competitiveness & Fee Policy) và điều chỉnh chiều tác động trong giả thuyết H1 và mô hình sang DƯƠNG (+)."
        ),
        (
            "3. Rà soát biến phụ thuộc DEC (Selection Decision):\nCần đo mức độ ưu tiên/lựa chọn VietinBank so với các ngân hàng khác để khớp tên đề tài.",
            "Bản cũ bị lẫn lộn giữa hài lòng chung (Satisfaction) và quyết định lựa chọn cạnh tranh.",
            "Chuẩn hóa 4 câu hỏi quan sát của DEC tập trung đo lường: Ưu tiên lựa chọn số 1, Phân bổ tỷ trọng bảo lãnh lớn tại VietinBank, Tiếp tục lựa chọn và Sẵn sàng giới thiệu đối tác."
        ),
        (
            "4. Thang đo phải bám sát Bảng hỏi gốc (n = 800):\nKhông tự xây dựng thang đo mới rồi gán ép vào bộ dữ liệu đã thu thập.",
            "Đảm bảo tính chân thực: Mọi biến trong đề cương phải có cột dữ liệu tương ứng trong file khảo sát 34 câu.",
            "Ánh xạ 100% hệ thống 7 biến độc lập và biến phụ thuộc khớp hoàn toàn với 34 câu hỏi thực tế trong Phiếu khảo sát VietinBank."
        ),
        (
            "5. Chuẩn hóa Mục tiêu & Câu hỏi nghiên cứu:\nViết lại mục tiêu tổng quát và 4 mục tiêu cụ thể tiếng Anh chuẩn NEU MDE, bổ sung phân tích khác biệt nhóm.",
            "Nâng cao chuẩn mực học thuật quốc tế và mở rộng góc nhìn quản trị thông qua kiểm định so sánh nhóm doanh nghiệp.",
            "Cập nhật nguyên văn 100% bộ Mục tiêu tiếng Anh của Giảng viên vào Mục 1.2 & 1.3; bổ sung kiểm định ANOVA & t-test ở Mục tiêu 3."
        )
    ]

    for fb, rsn, sol in fb_items:
        r_cells = t1.add_row().cells
        for i, w in enumerate(col_w1): r_cells[i].width = w
        r_cells[0].paragraphs[0].add_run(fb)
        r_cells[1].paragraphs[0].add_run(rsn)
        r_cells[2].paragraphs[0].add_run(sol)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # ==================== PHẦN II ====================
    doc.add_heading("PHẦN II: TỔNG HỢP CÁC NỘI DUNG ĐÃ ĐIỀU CHỈNH TRONG FILE THESIS DESIGN", level=1)
    
    add_body(
        "Toàn bộ các yêu cầu chỉnh sửa trên đã được cập nhật hoàn chỉnh trong file Đề cương Luận văn chính thức: "
        "DTL_Thesis_Design_NEU_MDE_Final.docx. Cụ thể từng chương mục như sau:"
    )

    doc.add_heading("2.1. Cập nhật Chương 1: Đặt vấn đề & Mục tiêu Nghiên cứu", level=2)
    add_bullet("Tiêu đề tiếng Anh: FACTORS AFFECTING CORPORATE CUSTOMERS’ DECISION TO CHOOSE BANK GUARANTEE SERVICES AT VIETNAM JOINT STOCK COMMERCIAL BANK FOR INDUSTRY AND TRADE (VIETINBANK).")
    add_bullet("Mục tiêu tổng quát (General Objective): Cập nhật nguyên văn chuẩn tiếng Anh cô giao: 'The general objective of this study is to identify and assess the factors associated with corporate customers’ decision to choose VietinBank for bank guarantee services, and to propose managerial recommendations for improving VietinBank’s attractiveness and competitiveness in the corporate bank guarantee market.'")
    add_bullet("4 Mục tiêu cụ thể (Specific Objectives 1-4): Cập nhật nguyên văn 4 mục tiêu tiếng Anh, trong đó Mục tiêu 3 tích hợp kiểm định khác biệt nhóm (Sub-group Analysis) theo loại hình doanh nghiệp, quy mô, thâm niên và loại bảo lãnh.")

    doc.add_heading("2.2. Cập nhật Chương 2: Khung Lý thuyết Nền tảng & Tổng quan", level=2)
    add_body("Bổ sung và phân định rõ ràng 5 khung lý thuyết bảo chứng cho mô hình 7 biến:")
    add_numbered("1.", "Lý thuyết Trung gian Tài chính & Giám sát Ủy thác (Diamond, 1984): Cơ sở của giá trị bảo chứng thương hiệu Big4.")
    add_numbered("2.", "Lý thuyết Định giá Rủi ro Tín dụng (Merton, 1974; Stiglitz & Weiss, 1981): Cơ sở của tính cạnh tranh biểu phí (COST_COMP) và ký quỹ (COLL_POLICY).")
    add_numbered("3.", "Mô hình SERVQUAL (Parasuraman et al., 1988): Cơ sở đo lường năng lực cán bộ (STAFF_QUAL), tốc độ (PROC_SPEED), và tách riêng Uy tín (BANK_REP).")
    add_numbered("4.", "Lý thuyết Ngân hàng Quan hệ (Boot, 2000; Berger & Udell, 1995): Cơ sở của biến Quan hệ tín dụng & Hạn mức (RELATIONSHIP).")
    add_numbered("5.", "Mô hình Chấp nhận Công nghệ TAM (Davis, 1989): Cơ sở giải thích động lực lựa chọn qua kênh số hóa eFAST (DIGITAL_CONV).")

    doc.add_heading("2.3. Cập nhật Chương 3: Mô hình Kinh tế lượng & Hệ thống Biến Thang đo", level=2)
    add_body("Phương trình Hồi quy Tuyến tính Đa biến (OLS) chính thức được thiết lập:")
    
    p_eq = doc.add_paragraph()
    p_eq.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_eq.paragraph_format.space_after = Pt(8)
    r_eq = p_eq.add_run(
        "DEC = β0 + β1*COST_COMP + β2*PROC_SPEED + β3*DIGITAL_CONV + β4*BANK_REP + β5*RELATIONSHIP + β6*STAFF_QUAL + β7*COLL_POLICY + ε"
    )
    r_eq.bold = True
    r_eq.font.size = Pt(10.5)

    add_body("Bảng hệ thống 7 biến độc lập và 1 biến phụ thuộc đã được khớp 100% với 34 câu hỏi trong bảng khảo sát gốc:")

    # Table 2: 7 Variables
    t2 = doc.add_table(rows=1, cols=5)
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.autofit = False
    col_w2 = [Inches(1.2), Inches(1.8), Inches(2.5), Inches(0.9), Inches(0.6)]

    hdr2 = t2.rows[0].cells
    hdr_titles2 = ["Mã biến", "Tên biến", "Nội dung thang đo (Khớp 34 câu hỏi gốc)", "Loại biến", "Dấu"]
    for idx, t in enumerate(hdr_titles2):
        hdr2[idx].width = col_w2[idx]
        p = hdr2[idx].paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(t)
        r.bold = True
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        set_cell_background(hdr2[idx], "003366")

    v_rows = [
        ("DEC", "Selection Decision", "Quyết định và mức độ ưu tiên chọn VietinBank so với đối thủ", "Phụ thuộc (Y)", "N/A"),
        ("COST_COMP", "Price Competitiveness", "Tính hợp lý, cạnh tranh của biểu phí và chiết khấu phí", "Độc lập (X1)", "+"),
        ("PROC_SPEED", "Processing Speed", "Thời gian thẩm định nhanh, thủ tục phát hành đơn giản", "Độc lập (X2)", "+"),
        ("DIGITAL_CONV", "Digital eFAST Convenience", "Nộp đề nghị online, cấp e-guarantee trực tuyến 24/7", "Độc lập (X3)", "+"),
        ("BANK_REP", "Bank Reputation", "Uy tín thương hiệu Big4, bảo lãnh được chấp nhận 100%", "Độc lập (X4)", "+"),
        ("RELATIONSHIP", "Relationship & Limits", "Hạn mức bảo lãnh linh hoạt, chính sách chăm sóc VIP", "Độc lập (X5)", "+"),
        ("STAFF_QUAL", "Staff Professionalism", "Trình độ cán bộ RM, am hiểu Luật Đấu thầu / TT 61", "Độc lập (X6)", "+"),
        ("COLL_POLICY", "Collateral Flexibility", "Tỷ lệ ký quỹ linh hoạt, đa dạng tài sản thế chấp", "Độc lập (X7)", "+")
    ]

    for code, name, desc, vtype, sign in v_rows:
        r_cells = t2.add_row().cells
        for i, w in enumerate(col_w2): r_cells[i].width = w
        r_cells[0].paragraphs[0].add_run(code).bold = True
        r_cells[1].paragraphs[0].add_run(name)
        r_cells[2].paragraphs[0].add_run(desc)
        p3 = r_cells[3].paragraphs[0]
        p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p3.add_run(vtype)
        p4 = r_cells[4].paragraphs[0]
        p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p4.add_run(sign).bold = True

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    doc.add_heading("2.4. Cập nhật Chương 4: Kết luận & Khuyến nghị Quản trị", level=2)
    add_body("Thiết lập 4 nhóm giải pháp thực tiễn kết nối trực tiếp với 4 mục tiêu nghiên cứu:")
    add_bullet("Giải pháp 1 (Theo Objective 1): Tập trung nguồn lực nâng cao 7 giá trị cốt lõi đã được định vị.")
    add_bullet("Giải pháp 2 (Theo Objective 2): Tối ưu hóa biểu phí cạnh tranh và rút ngắn thời gian xử lý phát hành bảo lãnh.")
    add_bullet("Giải pháp 3 (Theo Objective 3): Thiết kế chính sách sản phẩm may đo riêng theo phân khúc doanh nghiệp Nhà nước, Tư nhân/SME và FDI.")
    add_bullet("Giải pháp 4 (Theo Objective 4): Đẩy mạnh kênh số hóa VietinBank eFAST và tư vấn pháp lý chuyên sâu.")

    # Save output file
    output_filename = "c:/Users/nguyen.tuan.minh/Desktop/DTL-Master-Project/Bao_Cao_Tong_Hop_Gop_Y_Co_Giao_Huong_Dan.docx"
    doc.save(output_filename)
    print(f"Successfully generated Advisor Feedback Report at {output_filename}!")

if __name__ == "__main__":
    generate_report_docx()
